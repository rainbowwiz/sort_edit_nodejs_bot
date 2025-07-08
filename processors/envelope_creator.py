from docx import Document
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def normalize_name(name):
    return name.strip().lower().replace(" ", "_")

def find_docx(first, last, people_dirs):
    target = normalize_name(f"{first}_{last}")
    for d in people_dirs:
        for file in d.glob("*.docx"):
            if normalize_name(file.stem) == target:
                return file
    return None

def copy_doc_content(source_doc, target_doc):
    source_section = source_doc.sections[0]
    target_section = target_doc.sections[0]

    target_section.page_width = source_section.page_width
    target_section.page_height = source_section.page_height
    target_section.left_margin = source_section.left_margin
    target_section.right_margin = source_section.right_margin
    target_section.top_margin = source_section.top_margin
    target_section.bottom_margin = source_section.bottom_margin

    paragraphs = source_doc.paragraphs[:11]
    for para in paragraphs:
        new_para = target_doc.add_paragraph()
        new_para.style = para.style
        new_para.alignment = para.alignment
        pf = new_para.paragraph_format
        pf.left_indent = para.paragraph_format.left_indent
        pf.right_indent = para.paragraph_format.right_indent
        pf.first_line_indent = para.paragraph_format.first_line_indent
        pf.space_after = 0
        pf.space_before = 0

        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.name:
                new_run.font.name = run.font.name
                r = new_run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), run.font.name)
            new_run.font.size = run.font.size
            if run.font.color and run.font.color.rgb:
                new_run.font.color.rgb = run.font.color.rgb

def create_envelope_docs(combined_info, people_dirs, envelope_path: Path):
    envelope_path.mkdir(exist_ok=True)

    for item in combined_info:
        combined_file = item['pdf']
        name_list = item['names']
        doc = Document()

        for i, (last, first) in enumerate(name_list):
            doc_path = find_docx(first, last, people_dirs)
            if doc_path:
                person_doc = Document(str(doc_path))
                copy_doc_content(person_doc, doc)
                if i < len(name_list) - 1:
                    doc.add_page_break()
            else:
                print(f"⚠️ No docx found for: {first} {last}")

        output_filename = f"{combined_file.stem}.docx"
        doc.save(envelope_path / output_filename)
        print(f"📄 Saved envelope: {output_filename}")
